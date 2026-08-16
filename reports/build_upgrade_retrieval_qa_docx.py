"""Build the retrieval/query-parser/QA implementation memo."""
from __future__ import annotations

import json
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "data"
    / "reports"
    / "KE_HOACH_UPGRADE_RETRIEVAL_QA_QUERY_PARSER.docx"
)
BLUE, DARK_BLUE, INK = "2E74B5", "1F4D78", "0B2545"
MUTED, LIGHT_GRAY, CALLOUT, BORDER = "667085", "F2F4F7", "F4F6F9", "D0D5DD"
GREEN, GOLD, RED = "1F5A44", "7A5A00", "9B1C1C"
CONTENT_DXA, TABLE_INDENT_DXA = 9360, 120


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style(style, name, size, color=None, before=0, after=6, line=1.10):
    style.font.name = name
    style.font.size = Pt(size)
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), name)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    fmt = style.paragraph_format
    fmt.space_before, fmt.space_after, fmt.line_spacing = Pt(before), Pt(after), line


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.right_margin = Inches(1)
    section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    set_style(doc.styles["Normal"], "Calibri", 11, after=6, line=1.10)
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        set_style(style, "Calibri", size, color, before, after, 1.0)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
    code = doc.styles.add_style("Code Block", 1)
    set_style(code, "Consolas", 8.5, INK, before=4, after=6, line=1.0)
    code.paragraph_format.left_indent = code.paragraph_format.right_indent = Inches(0.15)
    citation = doc.styles.add_style("Table Citation", 1)
    set_style(citation, "Calibri", 8.5, MUTED, before=4, after=4, line=1.0)
    citation.font.italic = True
    configure_header_footer(doc)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, placeholder, end])
    set_font(run, size=8.5, color=MUTED)


def configure_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), 2)
    set_font(header.add_run("RETRIEVAL / QA UPGRADE"), size=8.5, color=MUTED, bold=True)
    set_font(header.add_run("\tTechnical implementation memo"), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    set_font(footer.add_run("Trang "), size=8.5, color=MUTED)
    add_field(footer, "PAGE")


def shade(node, fill):
    props = node._tc.get_or_add_tcPr() if hasattr(node, "_tc") else node._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    props.append(shd)


def bottom_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for key, value in (("val", "single"), ("sz", "16"), ("space", "1"), ("color", BLUE)):
        bottom.set(qn(f"w:{key}"), value)
    borders.append(bottom)
    p_pr.append(borders)


def create_numbering(doc, bullet):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id, num_id = max(abstract_ids, default=-1) + 1, max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"A1B2{abstract_id:04X}"[-8:])
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    template = OxmlElement("w:tmpl")
    template.set(qn("w:val"), f"C3D4{abstract_id:04X}"[-8:])
    abstract.append(template)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, value in (
        ("start", "1"),
        ("numFmt", "bullet" if bullet else "decimal"),
        ("lvlText", "\uf0b7" if bullet else "%1."),
        ("suff", "tab"),
    ):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:val"), value)
        lvl.append(node)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    lvl.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    for key, value in (("before", "0"), ("after", "160"), ("line", "280"), ("lineRule", "auto")):
        spacing.set(qn(f"w:{key}"), value)
    p_pr.append(spacing)
    lvl.append(p_pr)
    if bullet:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Symbol")
        r_fonts.set(qn("w:hAnsi"), "Symbol")
        r_fonts.set(qn("w:hint"), "default")
        r_pr.append(r_fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    # OOXML requires every abstractNum before the first concrete num.  Word
    # silently renders an out-of-order custom bullet as decimal numbering.
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def add_list(doc, text, num_id):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    ref = OxmlElement("w:numId")
    ref.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, ref])
    p._p.get_or_add_pPr().append(num_pr)
    set_font(p.add_run(text))


def set_cell(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")
    mar = OxmlElement("w:tcMar")
    for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_geometry(table, widths):
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"invalid table widths: {widths}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag, width in (("tblW", CONTENT_DXA), ("tblInd", TABLE_INDENT_DXA)):
        node = tbl_pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_pr.append(node)
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        for key, value in (("val", "single"), ("sz", "4"), ("space", "0"), ("color", BORDER)):
            node.set(qn(f"w:{key}"), value)
        borders.append(node)
    tbl_pr.append(borders)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            set_cell(cell, widths[index])


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    set_table_geometry(table, widths)
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(repeat)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0:
                shade(cell, LIGHT_GRAY)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_font(run, size=9, color=INK, bold=row_index == 0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(doc, title, body, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(title), size=10.5, color=color, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    set_font(p.add_run(body), size=10)
    shade(cell, CALLOUT)
    set_table_geometry(table, [CONTENT_DXA])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True


def body(doc, text):
    p = doc.add_paragraph()
    set_font(p.add_run(text))


def code(doc, text):
    p = doc.add_paragraph(style="Code Block")
    shade(p, CALLOUT)
    p.paragraph_format.keep_together = True
    set_font(p.add_run(text), name="Consolas", size=8.5, color=INK)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def masthead(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run("KẾ HOẠCH & IMPLEMENTATION MEMO"), size=23, color="000000", bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    set_font(
        p.add_run("Nâng cấp Retrieval, Typed Query Parser và Evidence-Grounded QA — không ASR"),
        size=14,
        color="373737",
    )
    for label, value in (
        ("Branch", "feat/update-new-model"),
        ("Baseline", "240b1b851be17a22b7b39e00117d052ce3ca5ac0"),
        ("Phạm vi", "Backend retrieval, QA evaluation, model integration và DOCX"),
        ("Ngoài phạm vi", "ASR, query expansion ownership, temporal implementation, frontend"),
        ("Ngày kiểm tra", "13/08/2026"),
        ("Trạng thái", "Đã implement; quality/model gates cần artifacts thật"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(f"{label}: "), size=10.5, bold=True)
        set_font(p.add_run(value), size=10.5)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    bottom_rule(p)


def executive_summary(doc, bullets):
    heading(doc, "Quyết định điều hành")
    add_callout(
        doc,
        "Kết luận",
        "Giữ nguyên offline feature pipeline và đường search KIS/AVS hiện hữu. Bổ sung một "
        "QA path có contract bất biến, routing có trace, evidence có provenance, Qwen answerer "
        "và hai tầng BGE. Mọi model mới đều lazy-load và bị khóa sau feature flag.",
    )
    for text in (
        "Không khôi phục Whisper/ASR; caption, OCR, object, BGE index và QA prompt mới không đọc transcript.",
        "QA/parser không sinh query expansion. expanded_queries chỉ được passthrough từ owner bên ngoài.",
        "Parser chỉ phát needs_temporal=true; không gọi hoặc thay đổi temporal retrieval.",
        "Không sửa frontend và không thay SigLIP2, Qwen caption, PP-OCRv5, YOLOE, keyframe selection hay RRF hiện có.",
        "Không công bố tăng chất lượng cho đến khi có data/public, labels thật, baseline artifacts và real-checkpoint run.",
    ):
        add_list(doc, text, bullets)
    page_break(doc)
    heading(doc, "Tình trạng triển khai theo phase", 2)
    add_table(
        doc,
        ["Phase", "Đã có trong branch", "Gate hiện tại"],
        [
            ("0 — Eval", "Fixture song ngữ, metrics, split/receipt", "Code pass; quality pending dữ liệu thật"),
            ("1 — Parser", "Typed QueryPlan bất biến, VI/EN", "Unit/regression pass"),
            ("2 — Router", "Hints, weighted RRF k=60, trace", "Unit pass; Hit@10 pending"),
            ("3 — Evidence", "Top-20, same-shot dedupe, Top-5 bundle", "Unit pass; provenance corpus pending"),
            ("4 — Qwen QA", "Top-3 answer, cache/timeout/fallback", "Contract pass; checkpoint pending"),
            ("5 — BGE-M3", "Dense 1024-d, IndexFlatIP, lineage", "Artifact tests pass; index chưa build"),
            ("6 — Reranker", "Top-100 → Top-20, alpha/fallback", "Runner tests pass; latency pending"),
        ],
        [1180, 4100, 4080],
    )


def assessment(doc):
    heading(doc, "1. Hiện trạng và đối chiếu nhận định từ đoạn chat")
    body(
        doc,
        "Nhận định tổng quát “feature extraction đã mạnh nhưng retrieval/ranking chưa theo kịp” "
        "là đúng về hướng ưu tiên. Một số dòng trong bảng chat đã cũ so với baseline 240b1b85 "
        "hoặc không còn phù hợp với quyết định scope của nhóm.",
    )
    add_table(
        doc,
        ["Nhận định trong chat", "Đối chiếu baseline", "Quyết định"],
        [
            ("SigLIP2 So400m đủ mạnh", "Đúng; visual contract và FAISS đã ổn định", "Giữ, chỉ benchmark"),
            ("Qwen3-VL caption là upgrade lớn", "Caption hiện dùng Qwen3-VL-8B-Instruct revision đã pin", "Giữ caption path"),
            ("PP-OCRv5 nên giữ", "Đúng; detector + Latin recognizer hiện hữu", "Giữ BM25 exact OCR"),
            ("YOLOE mask có thể nặng", "Hợp lý nhưng chưa có bằng chứng thay model", "Không đổi trong scope"),
            ("ASR là P0 bắt buộc", "Đúng cho speech query, trái quyết định resource hiện tại", "Loại khỏi vòng test"),
            ("Dense text retrieval còn thiếu", "Đúng tại baseline", "Thêm BGE-M3 dense-only"),
            ("Fusion chỉ heuristic", "Không còn đúng; weighted RRF k=60 đã tồn tại", "Tái sử dụng"),
            ("Text reranker còn thiếu", "Đúng tại baseline", "Thêm bge-reranker-v2-m3"),
            ("Keyframe selection chưa explicit", "Không đúng; selection/coverage đã có", "Không sửa"),
            ("Query parser chưa có", "QueryPlan có, nhưng QA có parser rời/expansion nội bộ", "Hợp nhất parser v2"),
            ("Temporal retrieval chưa có", "Không đúng; temporal path đã tồn tại", "Chỉ handoff"),
            ("SmolVLM2 có thể thừa", "Là VLM rerank tùy chọn, khác vai trò QA", "Không xóa"),
        ],
        [2600, 3380, 3380],
    )
    add_callout(
        doc,
        "Điểm cần hiểu đúng về ASR",
        "Bỏ ASR làm giảm recall cho câu hỏi dựa trên lời nói. Đây là trade-off có chủ đích để "
        "giảm tải GPU trong vòng test hiện tại, không phải kết luận ASR vô ích.",
        GOLD,
    )


def scope(doc, bullets):
    heading(doc, "2. Phạm vi giữ nguyên, nâng cấp và không làm")
    add_table(
        doc,
        ["Nhóm", "Thành phần"],
        [
            (
                "Giữ nguyên",
                "TransNetV2; dense candidates; keyframe selection/event coverage; SigLIP2; "
                "Qwen caption; PP-OCRv5; YOLOE; FAISS visual; BM25; weighted RRF; temporal; KIS/AVS.",
            ),
            (
                "Nâng cấp",
                "Typed QueryPlan; QA router; evidence bundle; Qwen grounded answer; BGE-M3 "
                "dense text; BGE reranker; QA evaluator; /retrieval/qa.",
            ),
            (
                "Không làm",
                "Whisper/ASR; tự dịch/synonym/query expansion; temporal neighbour expansion; "
                "đổi extractor; viết lại RRF/keyframe coverage; frontend.",
            ),
        ],
        [1800, 7560],
    )
    heading(doc, "Ownership boundary", 2)
    for text in (
        "QA/parser owner: typed intent, constraints, modality hints, evidence và answer contract.",
        "Expansion owner: cung cấp expanded_queries; QA không chỉnh hoặc tự bổ sung.",
        "Temporal owner: đọc needs_temporal/temporal_events; QA không gọi implementation.",
        "Offline owner: duy trì metadata/index; Phase 5 chỉ đọc metadata để build index mới.",
        "Frontend owner: không có file frontend nào thuộc patch này.",
    ):
        add_list(doc, text, bullets)


def pipelines(doc):
    heading(doc, "3. Pipeline hiện tại, rollout và pipeline đích")
    heading(doc, "3.1 Offline pipeline giữ nguyên", 2)
    code(
        doc,
        "Video\n"
        "  → TransNetV2 shot boundaries\n"
        "  → Dense candidates + keyframe selection / coverage\n"
        "  → SigLIP2 embeddings + Qwen caption + PP-OCRv5 + YOLOE\n"
        "  → FAISS visual + metadata + BM25 caption/OCR/object",
    )
    heading(doc, "3.2 Trình tự rollout theo phase", 2)
    code(
        doc,
        "Baseline/Eval → Typed Parser → QA Router → Evidence Bundle\n"
        "              → Qwen Grounded Answer → BGE-M3 Dense → BGE Reranker",
    )
    heading(doc, "3.3 Runtime đích sau khi bật đủ Phase 1–6", 2)
    code(
        doc,
        "Question + task_mode + external expanded_queries\n"
        "  → Typed QueryPlan (answer slot, constraints, hints, temporal handoff)\n"
        "  → parallel: SigLIP2 visual | BM25 caption/OCR/objects | BGE-M3 dense\n"
        "  → weighted RRF (k=60; hinted modality ×1.5; ≤100/modality)\n"
        "  → BGE reranker (Top 100 → Top 20; alpha=0.5)\n"
        "  → same-shot dedupe → Top 5 Evidence Bundle\n"
        "  → Qwen3.5 reads Top 3 → answered | insufficient_evidence",
    )
    add_callout(
        doc,
        "Hai nhánh chỉ bàn giao, không thực thi",
        "expanded_queries được search nguyên văn và ghi source=external_expansion. "
        "needs_temporal=true chỉ xuất hiện trong plan/trace; temporal engine không được gọi.",
    )


def phases(doc, bullets):
    heading(doc, "4. Chi tiết Phase 0 và sáu phase nâng cấp")
    items = [
        (
            "Phase 0 — Baseline & QA evaluation",
            [
                "Fixture 18 record synthetic: dev VI + locked-test EN; đủ object, color, OCR, action, count, location, yes/no, identity và unanswerable.",
                "Metrics: Hit@1/5/10, MRR, nDCG@10, EM/token-F1, abstention, parser accuracy/F1, latency và VRAM.",
                "Locked-test workflow tạo hash receipt atomic để chặn reuse âm thầm.",
                "quality_claim_allowed=false; gate thật pending data/public và adjudicated labels.",
            ],
        ),
        (
            "Phase 1 — Typed Query Parser v2",
            [
                "QueryPlan là frozen contract duy nhất; parser QA rời đã bị loại.",
                "task_mode explicit thắng inference; answer_type/constraints dùng rule VI/EN, không dùng LLM.",
                "QA không gọi expansion nội bộ; unknown giữ câu gốc và hybrid fallback.",
                "Before/after/then chỉ tạo temporal_relation, temporal_events và needs_temporal.",
                "Flag QA_TYPED_PARSER_ENABLED=true; rollback tạo plan tối giản.",
            ],
        ),
        (
            "Phase 2 — QA-aware Retrieval Router",
            [
                "OCR → OCR/caption; object/count → visual/objects/caption; color/action/location/identity → visual/caption.",
                "Mỗi modality tối đa 100 candidate; weighted RRF k=60; modality hint nhân 1.5.",
                "External expanded queries được fusion/truy vết nhưng không sinh tại parser.",
                "Flag QA_ROUTER_ENABLED=true; rollback gọi hybrid engine hiện hữu.",
            ],
        ),
        (
            "Phase 3 — Evidence Bundle",
            [
                "Top 20 pool khi chưa rerank; có reranker thì giữ Top 100 đầu vào và Top 20 đầu ra.",
                "Dedupe cùng shot rồi chọn Top 5; evidence có E001…, video/frame/shot/time/path, raw OCR, objects, modalities, score.",
                "Warning cho missing_image_path, missing_shot_metadata và conflicting_metadata.",
                "Không kéo temporal neighbours; /qa-evidence và mode=qa vẫn trả legacy results.",
                "Flag QA_EVIDENCE_BUNDLE_ENABLED=true; rollback giữ results.",
            ],
        ),
        (
            "Phase 4 — Qwen3.5 Grounded QA",
            [
                "Qwen/Qwen3.5-9B@c202236 đọc tối đa Top 3 evidence và chỉ trả JSON contract.",
                "Lazy shared runner, CUDA 4-bit mặc định, generation lock tuần tự, cache theo question/evidence/model/prompt revision.",
                "Citation phải thuộc evidence ID; answer type phải khớp parser; JSON/OOM/timeout fallback manual evidence.",
                "QA_ANSWER_MODE=off|optional|required, mặc định off.",
            ],
        ),
        (
            "Phase 5 — BGE-M3 Dense Text",
            [
                "BAAI/bge-m3 dense-only 1024-d, CLS pooling, float32 normalized, FAISS IndexFlatIP.",
                "Document chỉ có [CAPTION], [OCR], [OBJECTS]; legacy ASR key bị bỏ qua.",
                "Canonical ordering, source/document hashes, artifact SHA256 và resolved model revision.",
                "QA_BGE_DENSE_ENABLED=false mặc định; BM25 luôn giữ.",
            ],
        ),
        (
            "Phase 6 — BGE Cross-encoder Reranker",
            [
                "BAAI/bge-reranker-v2-m3 chấm tối đa Top 100, trả Top 20.",
                "score = alpha × retrieval + (1-alpha) × reranker; alpha mặc định 0.5.",
                "Candidate thiếu text giữ retrieval score; model/scoring error fallback Phase 5.",
                "QA_BGE_RERANKER_ENABLED=false; benchmark alpha 0.3/0.5/0.7 trên dev.",
            ],
        ),
    ]
    for index, (title, lines) in enumerate(items):
        heading(doc, title, 2)
        for line in lines:
            add_list(doc, line, bullets)


def api_contract(doc):
    heading(doc, "5. API, schema và compatibility")
    add_table(
        doc,
        ["Interface", "Hành vi"],
        [
            ("/retrieval/qa-evidence", "Evidence thủ công + legacy results"),
            ("mode=qa", "Alias cũ, không gọi answerer"),
            ("/retrieval/qa", "Typed plan + trace + answer + evidence + latency"),
            ("mode=qa_answer", "Alias end-to-end của /retrieval/qa"),
        ],
        [2600, 6760],
    )
    heading(doc, "Request", 2)
    code(
        doc,
        '{\n'
        '  "query": "Người phụ nữ áo đỏ cầm gì?",\n'
        '  "top_k": 5,\n'
        '  "task_mode": "qa",\n'
        '  "expanded_queries": []\n'
        "}",
    )
    heading(doc, "Response data", 2)
    code(
        doc,
        '{\n'
        '  "query_plan": {\n'
        '    "task_mode":"qa", "answer_type":"object",\n'
        '    "retrieval_statement":"Người phụ nữ áo đỏ cầm một vật",\n'
        '    "known_constraints":{"subject":["người phụ nữ"],"attributes":["áo đỏ"],"actions":["cầm"]},\n'
        '    "modality_hints":["objects","visual","caption"], "needs_temporal":false, "confidence":0.94\n'
        "  },\n"
        '  "routing_trace":{"rrf_k":60,"hint_boost":1.5,"queries":[]},\n'
        '  "answer":{"status":"answered","answer":"một chiếc điện thoại","answer_type":"object","confidence":0.86,"evidence_ids":["E001"]},\n'
        '  "evidence":[{"evidence_id":"E001","video_id":"...","frame_id":"..."}],\n'
        '  "latency_ms":0.0\n'
        "}",
    )
    body(
        doc,
        "FastAPI vẫn dùng envelope success/data/message như các retrieval endpoint hiện hữu; "
        "schema trên là nội dung trường data. Nếu mode=required thất bại, API trả HTTP 503 với "
        "success=false và vẫn giữ evidence trong data để fallback thủ công.",
    )


def tests_and_gates(doc):
    heading(doc, "6. Test matrix, acceptance gates và rollback")
    heading(doc, "6.1 Verification đã chạy", 2)
    add_table(
        doc,
        ["Hạng mục", "Kết quả tại branch"],
        [
            ("Focused QA/BGE/evaluator", "58 tests pass sau integration và rollback audit"),
            ("Backend full suite", "167 pass, 1 skip"),
            ("Competition full suite", "74 tests pass"),
            ("GPU runtime", "PyTorch 2.13.0+cu130; RTX 3050 Laptop 6 GB; CUDA=true"),
            ("Real checkpoints", "Chưa chạy: Qwen/BGE cache và BGE artifacts không tồn tại"),
            ("Quality benchmark", "Chưa chạy: data/public và adjudicated labels không tồn tại"),
        ],
        [3000, 6360],
    )
    add_callout(
        doc,
        "Không được diễn giải sai",
        "Unit tests chứng minh contract, fallback, lineage và integration; không chứng minh "
        "Evidence Hit, answer F1, hallucination, latency P95 hoặc VRAM checkpoint thật.",
        RED,
    )
    heading(doc, "6.2 Acceptance gates trước promote", 2)
    add_table(
        doc,
        ["Phase", "Gate bắt buộc", "Hiện trạng"],
        [
            ("1", "Task-mode ≥95%; answer-type macro F1 ≥90%; constraint F1 ≥90%; no regression", "Regression pass; metrics pending"),
            ("2", "Hit@10 +≥3 pp; OCR không giảm; trace đầy đủ", "Trace pass; quality pending"),
            ("3", "Hit@5 không giảm; no duplicate; provenance hợp lệ", "Unit pass; corpus pending"),
            ("4", "Citation ≥95%; hallucination ≤2%; no OOM; EM/F1 tăng", "Contract pass; model pending"),
            ("5", "Recall/nDCG tăng; Recall@100 giảm ≤0.01; temporal/VKIS giữ", "Lineage pass; benchmark pending"),
            ("6", "nDCG/F1 tăng; P95 +≤25%; OCR exact không giảm", "Fallback pass; benchmark pending"),
        ],
        [900, 5850, 2610],
    )
    heading(doc, "6.3 Rollback map", 2)
    add_table(
        doc,
        ["Biến", "Default", "Rollback effect"],
        [
            ("QA_TYPED_PARSER_ENABLED", "true", "Plan tối giản, original query"),
            ("QA_ROUTER_ENABLED", "true", "Existing hybrid search"),
            ("QA_EVIDENCE_BUNDLE_ENABLED", "true", "Legacy results còn giữ"),
            ("QA_ANSWER_MODE", "off", "Không load Qwen answerer"),
            ("QA_BGE_DENSE_ENABLED", "false", "Không load BGE index/model"),
            ("QA_BGE_RERANKER_ENABLED", "false", "Bỏ cross-encoder"),
            ("QA_BGE_RERANKER_ALPHA", "0.5", "Đổi alpha đã khóa từ dev"),
        ],
        [3450, 1300, 4610],
    )


def dependencies(doc, bullets):
    page_break(doc)
    heading(doc, "7. Dependency, model sources và artifact impact")
    add_table(
        doc,
        ["Vai trò", "Model / revision", "Artifact hoặc runtime impact"],
        [
            ("Visual", "google/siglip2-so400m-patch16-384", "Giữ visual artifacts"),
            ("Caption + QA", "Caption Qwen3-VL-8B-Instruct; grounded QA giữ model riêng", "QA cache/model path riêng"),
            ("OCR", "PP-OCRv5 server det + Latin mobile rec", "Giữ OCR JSON/BM25"),
            ("Object", "YOLOE-26l", "Giữ object metadata/BM25"),
            ("Dense text", "BAAI/bge-m3 @ resolved revision", "Thêm 3 BGE artifacts"),
            ("Reranker", "BAAI/bge-reranker-v2-m3", "Online model cache"),
            ("Legacy optional", "HuggingFaceTB/SmolVLM2-2.2B-Instruct", "Không đổi"),
        ],
        [1500, 3640, 4220],
    )
    heading(doc, "Requirements", 2)
    for text in (
        "requirements.txt là manifest chuẩn cho dependency dùng chung, backend, competition và script báo cáo.",
        "PyTorch và PaddlePaddle được cài riêng theo đúng một profile CPU/CUDA để tránh wheel xung đột.",
        "BGE dùng Transformers/FAISS sẵn có, không cần FlagEmbedding hoặc sentence-transformers.",
        "Không có faster-whisper, ctranslate2 hoặc dependency ASR mới.",
    ):
        add_list(doc, text, bullets)
    heading(doc, "BGE artifacts mới", 2)
    code(
        doc,
        "data/indexes/bge_m3/\n"
        "  ├─ bge_m3_flat_ip.faiss\n"
        "  ├─ bge_m3_frame_map.json\n"
        "  └─ bge_m3_manifest.json",
    )
    body(
        doc,
        "Manifest khóa dense/1024/normalized float32/IndexFlatIP, canonical ordering, "
        "resolved revision, source/document digests và SHA256 của index/frame map.",
    )
    heading(doc, "Lệnh build index từ metadata hiện có", 2)
    code(
        doc,
        "python -m backend.app.services.indexing.build_bge_m3_index "
        "--metadata data/metadata --output-root data/indexes/bge_m3 "
        "--model-revision <PINNED_COMMIT> --device cuda",
    )
    heading(doc, "Cấu hình rollout an toàn", 2)
    code(
        doc,
        '$env:QA_EXPERIMENT_ID="qa-bge-dev-001"\n'
        '$env:QA_ANSWER_MODE="off"\n'
        '$env:QA_BGE_DENSE_ENABLED="true"\n'
        '$env:QA_BGE_INDEX_ROOT="data/indexes/bge_m3"\n'
        '$env:QA_BGE_RERANKER_ENABLED="false"',
    )


def rollout(doc, decimal):
    heading(doc, "8. Thứ tự rollout được phép")
    for text in (
        "Freeze real dev labels và tạo baseline run tại 240b1b85; không mở locked test.",
        "Bật parser/router/evidence, chạy parser metrics và Evidence Hit/nDCG theo answer type.",
        "Chỉ khi Phase 1–3 qua gate mới bật QA_ANSWER_MODE=optional và smoke Qwen.",
        "Build BGE-M3 từ metadata, xác minh manifest/tamper rồi benchmark dense off/on.",
        "Benchmark alpha 0.3/0.5/0.7 trên dev, khóa alpha/revision trước locked test.",
        "Chạy GPU smoke, P50/P95 latency, peak VRAM/OOM và quality suite; promote từng flag.",
        "Locked test chạy một lần bằng receipt; gate fail thì rollback flag, không sửa label.",
    ):
        add_list(doc, text, decimal)
    add_callout(
        doc,
        "Điều kiện bàn giao cuối",
        "“Đã implement” không đồng nghĩa “đã promote”. Qwen/BGE vẫn mặc định off cho đến khi "
        "có artifact, checkpoint smoke và benchmark thật.",
        GREEN,
    )


def audit(doc):
    section = doc.sections[0]
    assert section.page_width == Inches(8.5)
    assert section.page_height == Inches(11)
    assert section.left_margin == section.right_margin == Inches(1)
    assert section.top_margin == section.bottom_margin == Inches(1)
    assert abs(section.header_distance.inches - 0.492) < (1 / 1440)
    assert abs(section.footer_distance.inches - 0.492) < (1 / 1440)
    for name, size, before, after in (
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = doc.styles[name]
        assert style.font.size == Pt(size)
        assert style.paragraph_format.space_before == Pt(before)
        assert style.paragraph_format.space_after == Pt(after)
    for table in doc.tables:
        widths = [
            int(node.get(qn("w:w")))
            for node in table._tbl.tblGrid.findall(qn("w:gridCol"))
        ]
        assert sum(widths) == CONTENT_DXA
        props = table._tbl.tblPr
        assert props.find(qn("w:tblW")).get(qn("w:w")) == str(CONTENT_DXA)
        assert props.find(qn("w:tblInd")).get(qn("w:w")) == str(TABLE_INDENT_DXA)
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                assert tc_w is not None and int(tc_w.get(qn("w:w"))) == widths[index]
    return {
        "preset": "standard_business_brief",
        "header_template": "memo_masthead",
        "page": "Letter portrait",
        "content_width_dxa": CONTENT_DXA,
        "table_count": len(doc.tables),
        "paragraph_count": len(doc.paragraphs),
    }


def build():
    doc = Document()
    doc.core_properties.title = "Kế hoạch nâng cấp Retrieval, QA và Query Parser"
    doc.core_properties.subject = "Implementation memo — no ASR"
    doc.core_properties.author = "AIChallenge26 Retrieval Team"
    configure_document(doc)
    bullets = create_numbering(doc, True)
    decimal = create_numbering(doc, False)
    masthead(doc)
    executive_summary(doc, bullets)
    assessment(doc)
    scope(doc, bullets)
    pipelines(doc)
    phases(doc, bullets)
    api_contract(doc)
    tests_and_gates(doc)
    dependencies(doc, bullets)
    rollout(doc, decimal)
    report = audit(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    report["output"] = str(OUTPUT)
    report["bytes"] = OUTPUT.stat().st_size
    return report


if __name__ == "__main__":
    print(json.dumps(build(), ensure_ascii=False, indent=2))
